"""Regression tests for the Word renderer used by protocol skeletons."""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from render_copyable_docx import PENDING_RED, build  # noqa: E402
from render_protocol_docx import FOOTER  # noqa: E402
from validate_protocol_docx import validate  # noqa: E402


class ProtocolDocxRendererTest(unittest.TestCase):
    def test_bilingual_protocol_uses_real_numbering_and_marks_both_pending_languages(self) -> None:
        markdown = """# 中国通用研究计划书骨架 | China-General Research Protocol Skeleton

1. 每个待填事实必须确认。| Every fact placeholder must be confirmed.

## 1. 研究基本信息与治理 | Core Study Information and Governance

### 1. 研究题目 | Study Title (`P-01`)

- 同源事实组 | Shared fact group: `P-01`
- 中文待填事实: `[由研究者／用户确认后填写]`
- English paired fact: `[To be completed after researcher/user confirmation]`

[待用户确认：叙述性正文中的缺失事实也必须以红色显示。]
"""
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "protocol.docx"
            build(markdown, output, footer_text=FOOTER)
            document = Document(output)

        numbered = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("每个待填事实"))
        self.assertEqual("List Number", numbered.style.name)
        chinese_pending = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("中文待填事实"))
        english_pending = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("English paired fact"))
        narrative_pending = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("[待用户确认"))
        self.assertTrue(any(str(run.font.color.rgb) == str(PENDING_RED) for run in chinese_pending.runs))
        self.assertTrue(any(str(run.font.color.rgb) == str(PENDING_RED) for run in english_pending.runs))
        self.assertTrue(any(str(run.font.color.rgb) == str(PENDING_RED) for run in narrative_pending.runs))
        self.assertIn(FOOTER, document.sections[0].footer.paragraphs[0].text)

    def test_renderer_removes_session_metadata_and_validator_accepts_generated_file(self) -> None:
        markdown = """# 中国通用研究计划书骨架

由 `protocol-coverage-matrix.yaml` 生成。

## 1. 研究基本信息与治理

### 1. 研究题目 (`P-01`)

- 中文待填事实: `[由研究者／用户确认后填写]`
"""
        with tempfile.TemporaryDirectory() as directory:
            base = Path(directory)
            markdown_path = base / "protocol.md"
            output = base / "protocol.docx"
            markdown_path.write_text(markdown, encoding="utf-8")
            build(markdown, output, footer_text=FOOTER)
            with ZipFile(output) as package:
                self.assertNotIn("docProps/custom.xml", package.namelist())
                self.assertFalse(
                    any(
                        b"rsid" in package.read(name)
                        for name in package.namelist()
                        if name.startswith("word/") and name.endswith(".xml")
                    )
                )
            self.assertEqual([], validate(markdown_path, output))


if __name__ == "__main__":
    unittest.main()
