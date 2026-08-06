#!/usr/bin/env python3
"""Convert a generated Markdown filling checklist into a Word document.

Use the compact_reference_guide preset with one named override: the entire
checklist is monochrome, so it reads as a practical clinical-workflow form.
The input must already come from render_copyable_checklist.py.
"""

from __future__ import annotations

import argparse
import re
import sys
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_ASCII = "Calibri"
FONT_CJK = "Microsoft YaHei"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
SUGGESTION_BLUE = RGBColor(46, 117, 182)
PENDING_RED = RGBColor(192, 0, 0)
DEFAULT_FOOTER = "research-ethics V1｜请以平台当前可见字段为准"
PENDING_MARKERS = (
    "待用户确认",
    "由研究者／用户确认后填写",
    "[To be completed after researcher/user confirmation]",
)


def scrub_privacy_metadata(output: Path) -> None:
    """Remove machine/session metadata from a generated DOCX in place.

    A new python-docx document inherits Word's default rsid attributes even
    after its visible core properties are cleared.  They are not research
    facts, but public examples and shareable filling drafts should not retain
    avoidable document-session identifiers.  This small OOXML pass removes
    rsid attributes, custom properties, and their package references.
    """
    with ZipFile(output, "r") as source, tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, dir=output.parent
    ) as temporary:
        temporary_path = Path(temporary.name)
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as target:
            for name in source.namelist():
                if name == "docProps/custom.xml":
                    continue
                payload = source.read(name)
                if name.endswith(".xml"):
                    # Preserve the original namespace prefixes and mc:Ignorable
                    # declarations. Re-serializing Word OOXML with ElementTree
                    # can rename prefixes while leaving their string references
                    # behind, which makes Microsoft Word flag a healthy file as
                    # damaged. These narrowly scoped substitutions avoid that.
                    payload = re.sub(rb'\s+w:rsid[A-Za-z]*="[^"]*"', b"", payload)
                    payload = re.sub(rb"<w:rsids\b[^>]*>.*?</w:rsids>", b"", payload, flags=re.DOTALL)
                    payload = re.sub(rb"<w:rsid(?:Root)?\b[^>]*/>", b"", payload)
                    if name == "_rels/.rels":
                        payload = re.sub(
                            rb'<Relationship\b[^>]*Type="[^"]*/custom-properties"[^>]*/>',
                            b"",
                            payload,
                        )
                    elif name == "[Content_Types].xml":
                        payload = re.sub(
                            rb'<Override\b[^>]*PartName="/docProps/custom\.xml"[^>]*/>',
                            b"",
                            payload,
                        )
                target.writestr(name, payload)
    temporary_path.replace(output)


def set_font(run, *, size: float, bold: bool = False, color=BLACK) -> None:
    run.font.name = FONT_ASCII
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def plain_markdown(text: str) -> str:
    """Keep the document copyable without leaking Markdown control marks."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


def add_colored_fragment(paragraph, text: str, *, size: float = 11, color=BLACK) -> None:
    """Add text while keeping unresolved placeholders visibly distinct."""
    marker_pattern = re.compile("|".join(re.escape(marker) for marker in PENDING_MARKERS))
    previous_end = 0
    for match in marker_pattern.finditer(text):
        if match.start() > previous_end:
            run = paragraph.add_run(text[previous_end : match.start()])
            set_font(run, size=size, color=color)
        run = paragraph.add_run(match.group(0))
        set_font(run, size=size, color=PENDING_RED)
        previous_end = match.end()
    if previous_end < len(text):
        run = paragraph.add_run(text[previous_end:])
        set_font(run, size=size, color=color)


def configure(document: Document, *, footer_text: str = DEFAULT_FOOTER) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (12, 10, 5),
    }
    for name, (size, before, after) in tokens.items():
        style = document.styles[name]
        style.font.name = FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    bullet.font.name = FONT_ASCII
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    bullet.font.size = Pt(11)
    bullet.font.color.rgb = BLACK
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    numbered = document.styles["List Number"]
    numbered.font.name = FONT_ASCII
    numbered._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    numbered.font.size = Pt(11)
    numbered.font.color.rgb = BLACK
    numbered.paragraph_format.space_after = Pt(4)
    numbered.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(footer_text)
    set_font(run, size=8.5, color=MUTED)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_font(run, size=18, bold=True)


def add_heading(document: Document, level: int, text: str) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    text = plain_markdown(text)
    prefix = "建议填写／选择："
    if text.startswith(prefix):
        label = paragraph.add_run(prefix)
        set_font(label, size=11, color=BLACK)
        value_and_options = text[len(prefix) :]
        value, separator, options = value_and_options.partition("；可选：")
        add_colored_fragment(paragraph, value, color=SUGGESTION_BLUE)
        if separator:
            options_run = paragraph.add_run(separator + options)
            set_font(options_run, size=11, color=BLACK)
        return
    add_colored_fragment(paragraph, text, color=BLACK)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    add_colored_fragment(paragraph, plain_markdown(text), color=BLACK)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    # Pending facts can occur in an ordinary narrative paragraph as well as a
    # checklist bullet.  Keep the visible unresolved marker red in both forms.
    add_colored_fragment(paragraph, plain_markdown(text), color=BLACK)


def build(markdown: str, output: Path, *, footer_text: str = DEFAULT_FOOTER) -> None:
    document = Document()
    configure(document, footer_text=footer_text)
    properties = document.core_properties
    for name in ("author", "last_modified_by", "title", "subject", "keywords", "comments", "category"):
        setattr(properties, name, "")
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            add_title(document, line[2:])
        elif line.startswith("## "):
            add_heading(document, 1, line[3:])
        elif line.startswith("### "):
            add_heading(document, 2, line[4:])
        elif line.startswith("#### "):
            add_heading(document, 3, line[5:])
        elif line.startswith("##### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(plain_markdown(line[6:]))
            set_font(run, size=11, bold=True)
        elif line.startswith("- "):
            add_bullet(document, line[2:])
        elif re.match(r"^\d+\.\s+", line):
            add_numbered(document, re.sub(r"^\d+\.\s+", "", line))
        else:
            add_body(document, line)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    scrub_privacy_metadata(output)


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    build(args.markdown.read_text(encoding="utf-8"), args.output)
    print(f"Rendered {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
