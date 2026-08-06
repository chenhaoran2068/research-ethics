#!/usr/bin/env python3
"""Perform structural and privacy checks on a generated protocol-skeleton DOCX.

This check complements, but cannot replace, render-and-inspect visual QA.  It
is intentionally reusable so generated Chinese, English, and bilingual
protocol skeletons can be audited before sharing or publishing the skill.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from render_copyable_docx import PENDING_MARKERS, PENDING_RED
from render_protocol_docx import FOOTER


ID_PATTERN = re.compile(r"`([A-Z]-\d{2})`")
METADATA_FIELDS = (
    "author",
    "last_modified_by",
    "title",
    "subject",
    "keywords",
    "comments",
    "category",
)


def markdown_fact_ids(markdown: str) -> set[str]:
    """Return canonical module/chapter IDs emitted by the Markdown generator."""
    return set(ID_PATTERN.findall(markdown))


def _is_pending_red(paragraph) -> bool:
    for marker in PENDING_MARKERS:
        if marker in paragraph.text:
            return any(
                marker in run.text and str(run.font.color.rgb) == str(PENDING_RED)
                for run in paragraph.runs
            )
    return True


def validate(markdown_path: Path, document_path: Path) -> list[str]:
    """Return human-readable validation failures, or an empty list."""
    issues: list[str] = []
    markdown = markdown_path.read_text(encoding="utf-8")
    if "protocol-coverage-matrix.yaml" not in markdown:
        issues.append("Markdown does not identify the protocol coverage matrix")
    expected_ids = markdown_fact_ids(markdown)
    if not expected_ids:
        issues.append("Markdown contains no canonical fact/chapter identifiers")

    document = Document(document_path)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    footer_text = "\n".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.footer.paragraphs
    )
    missing_ids = sorted(identifier for identifier in expected_ids if identifier not in body_text)
    if missing_ids:
        issues.append("DOCX is missing canonical IDs: " + ", ".join(missing_ids))
    if FOOTER not in footer_text:
        issues.append("DOCX does not contain the protocol-skeleton footer")
    for paragraph in document.paragraphs:
        if not _is_pending_red(paragraph):
            issues.append(f"Pending marker is not red: {paragraph.text[:100]}")

    properties = document.core_properties
    populated = [name for name in METADATA_FIELDS if getattr(properties, name, "")]
    if populated:
        issues.append("DOCX core properties are populated: " + ", ".join(populated))

    with ZipFile(document_path) as package:
        names = package.namelist()
        if "docProps/custom.xml" in names:
            issues.append("DOCX contains custom document properties")
        rsid_files = [
            name
            for name in names
            if name.startswith("word/") and name.endswith(".xml") and b"rsid" in package.read(name)
        ]
        if rsid_files:
            issues.append("DOCX contains rsid session identifiers: " + ", ".join(rsid_files))
    return issues


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown", required=True, type=Path)
    parser.add_argument("--docx", required=True, type=Path)
    args = parser.parse_args()
    try:
        issues = validate(args.markdown, args.docx)
    except (OSError, UnicodeError, ValueError, KeyError) as exc:
        print(f"Protocol DOCX cannot be audited: {exc}", file=sys.stderr)
        return 2
    if issues:
        print(f"Protocol DOCX validation failed: {len(issues)} issue(s)", file=sys.stderr)
        for issue in issues:
            print(f"- {issue}", file=sys.stderr)
        return 1
    print("Protocol DOCX structural and privacy validation passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
