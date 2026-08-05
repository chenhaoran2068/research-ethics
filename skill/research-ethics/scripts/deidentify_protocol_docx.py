#!/usr/bin/env python3
"""Create a de-identified research-protocol DOCX without touching its source.

This is a conservative first-pass redactor for using a protocol as a skill
test input. It removes title/header/footer identity, labeled identity fields,
common Chinese institution/person references, dates, identifier-like strings,
phone numbers and email addresses.  Review the resulting copy before external
sharing: free text can still contain context-specific identifiers.
"""

from __future__ import annotations

import argparse
import re
import sys
from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


IDENTITY_LABELS = (
    "研究题目",
    "课题名称",
    "项目名称",
    "方案名称",
    "研究机构",
    "研究单位",
    "申办方",
    "项目负责人",
    "课题负责人",
    "主要研究者",
    "研究者姓名",
    "联系人",
    "联系电话",
    "电子邮箱",
    "E-mail",
    "单位地址",
    "伦理批件",
    "伦理编号",
    "项目编号",
    "注册号",
)
PERSONNEL_TABLE_MARKERS = (
    *IDENTITY_LABELS,
    "研究人员",
    "数据与分析",
    "本地队列数据",
    "学术指导",
    "承担科室",
)
LABEL_RE = re.compile(r"(" + "|".join(map(re.escape, IDENTITY_LABELS)) + r")\s*[：:]\s*.+")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?<!\d)(?:1[3-9]\d{9}|(?:0\d{2,3}[- ]?)?\d{7,8})(?!\d)")
DATE_RE = re.compile(r"(?<!\d)20\d{2}(?:[-/.年]\d{1,2}(?:[-/.月]\d{1,2}日?)?)?(?!\d)")
ID_RE = re.compile(r"(?<![A-Za-z0-9])[A-Za-z]{0,6}\d{6,}[A-Za-z0-9_\-]*(?![A-Za-z0-9])")
ORG_RE = re.compile(r"[\u4e00-\u9fff]{2,}(?:医院|大学|学院|研究所|公司)")
PERSON_TITLE_RE = re.compile(r"[\u4e00-\u9fff]{2,4}(?=(?:教授|主任医师|副主任医师|医师|研究员|博士))")
TITLE_STYLE_NAMES = {"Title", "Subtitle", "标题", "副标题"}


def iter_tables(tables: Iterable[Table]) -> Iterable[Table]:
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell.tables)


def set_text(paragraph: Paragraph, text: str) -> None:
    """Replace a whole paragraph while retaining its paragraph-level layout."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def scrub_inline(text: str, changes: Counter[str]) -> str:
    substitutions = (
        (EMAIL_RE, "[电子邮箱]", "email"),
        (PHONE_RE, "[联系电话]", "phone"),
        (DATE_RE, "[日期]", "date"),
        (ID_RE, "[编号]", "identifier"),
        (ORG_RE, "[研究机构]", "organization"),
        (PERSON_TITLE_RE, "[人员]", "person"),
    )
    result = text
    for pattern, replacement, key in substitutions:
        result, count = pattern.subn(replacement, result)
        changes[key] += count
    return result


def redact_paragraph(paragraph: Paragraph, changes: Counter[str]) -> None:
    text = paragraph.text
    if not text.strip():
        return
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if style_name in TITLE_STYLE_NAMES:
        set_text(paragraph, "[已脱敏研究方案]")
        changes["title_style"] += 1
        return
    label_match = LABEL_RE.search(text)
    if label_match:
        set_text(paragraph, f"{label_match.group(1)}：[已脱敏]")
        changes["labeled_field"] += 1
        return
    redacted = scrub_inline(text, changes)
    if redacted != text:
        set_text(paragraph, redacted)


def redact_table(table: Table, changes: Counter[str]) -> None:
    # A personnel/ownership table can have a labeled first row followed by
    # role-specific rows without identity labels.  Treat the whole table as
    # identity-sensitive once one such marker is present; otherwise names in
    # later rows could survive a row-by-row pass.
    table_is_identity_sensitive = any(
        marker in cell.text
        for row in table.rows
        for cell in row.cells
        for marker in PERSONNEL_TABLE_MARKERS
    )
    for row in table.rows:
        label_row = any(any(label in cell.text for label in IDENTITY_LABELS) for cell in row.cells)
        for index, cell in enumerate(row.cells):
            if (table_is_identity_sensitive and index > 0 and cell.text.strip()) or (
                label_row and index > 0 and cell.text.strip()
            ):
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        set_text(paragraph, "[已脱敏]")
                        changes["identity_table_value"] += 1
            else:
                for paragraph in cell.paragraphs:
                    redact_paragraph(paragraph, changes)
            for nested in cell.tables:
                redact_table(nested, changes)


def scrub_properties(document: Document, changes: Counter[str]) -> None:
    properties = document.core_properties
    for name in ("author", "last_modified_by", "title", "subject", "keywords", "comments", "category"):
        if getattr(properties, name, None):
            setattr(properties, name, "")
            changes["core_property"] += 1


def deidentify(source: Path, output: Path) -> Counter[str]:
    document = Document(source)
    changes: Counter[str] = Counter()
    first_body = next((paragraph for paragraph in document.paragraphs if paragraph.text.strip()), None)
    if first_body is not None:
        set_text(first_body, "[已脱敏研究方案]")
        changes["first_body_title"] += 1
    for paragraph in document.paragraphs:
        if paragraph is not first_body:
            redact_paragraph(paragraph, changes)
    for table in document.tables:
        redact_table(table, changes)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                set_text(paragraph, "[已脱敏页眉]")
                changes["header_footer"] += 1
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                set_text(paragraph, "[已脱敏页脚]")
                changes["header_footer"] += 1
        for table in section.header.tables:
            redact_table(table, changes)
        for table in section.footer.tables:
            redact_table(table, changes)
    scrub_properties(document, changes)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return changes


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    changes = deidentify(args.input, args.output)
    print("De-identified copy written:", args.output)
    print("Changes:", ", ".join(f"{key}={value}" for key, value in sorted(changes.items())) or "none")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
